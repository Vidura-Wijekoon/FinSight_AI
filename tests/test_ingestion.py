"""
Tests for FinSight AI — Ingestion Pipeline
"""
import io
import json
import tempfile
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from fastapi import HTTPException

from src.ingestion.chunker import DocumentChunker
from src.ingestion.text_extractor import TextExtractor


class TestTextExtractor:
    def setup_method(self):
        self.extractor = TextExtractor()

    def test_extract_txt(self):
        text = "Total Revenue: $1,250,000\nNet Profit: $320,000"
        result = self.extractor.extract(text.encode("utf-8"), "txt")
        assert "Revenue" in result
        assert "Profit" in result

    def test_extract_pdf(self):
        """Minimal PDF with extractable text."""
        import pypdf
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        # Blank page returns empty string — just verify it doesn't raise
        result = self.extractor.extract(buf.read(), "pdf")
        assert isinstance(result, str)

    def test_extract_csv(self):
        csv_data = b"Company,Revenue,EBITDA\nABC Corp,1000000,250000\nXYZ Inc,500000,120000"
        result = self.extractor.extract(csv_data, "csv")
        assert "ABC Corp" in result
        assert "1000000" in result

    def test_extract_docx(self):
        import docx
        doc = docx.Document()
        doc.add_paragraph("Annual Report FY2024")
        doc.add_paragraph("Revenue: $2.5B")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        result = self.extractor.extract(buf.read(), "docx")
        assert "Annual Report" in result
        assert "Revenue" in result

    def test_extract_xlsx(self):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Financials"
        ws.append(["Quarter", "Revenue", "Profit"])
        ws.append(["Q1 2024", 1200000, 300000])
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        result = self.extractor.extract(buf.read(), "xlsx")
        assert "Financials" in result
        assert "Revenue" in result

    def test_unsupported_type_raises(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            self.extractor.extract(b"data", "pptx")


class TestDocumentChunker:
    def setup_method(self):
        self.chunker = DocumentChunker(chunk_size=512, chunk_overlap=50)

    def test_chunk_basic_text(self):
        text = "Financial data. " * 100  # ~1600 chars
        chunks = self.chunker.chunk(text, {"doc_id": "test-001", "source_file": "test.txt"})
        assert len(chunks) > 1

    def test_chunk_metadata_populated(self):
        text = "This is a sample financial document with enough text. " * 20
        chunks = self.chunker.chunk(text, {"doc_id": "abc-123", "source_file": "report.pdf"})
        assert chunks[0].metadata["doc_id"] == "abc-123"
        assert chunks[0].metadata["source_file"] == "report.pdf"
        assert chunks[0].metadata["chunk_index"] == 0

    def test_chunk_indices_sequential(self):
        text = "sentence. " * 200
        chunks = self.chunker.chunk(text, {"doc_id": "x", "source_file": "x.txt"})
        indices = [c.metadata["chunk_index"] for c in chunks]
        assert indices == list(range(len(indices)))

    def test_empty_text_returns_empty_list(self):
        chunks = self.chunker.chunk("", {"doc_id": "x", "source_file": "x.txt"})
        assert chunks == []

    def test_chunk_size_respected(self):
        text = "X" * 2000
        chunks = self.chunker.chunk(text, {"doc_id": "x", "source_file": "x.txt"})
        for c in chunks:
            assert len(c.text) <= 512 + 50  # Slight tolerance for splitter behavior

    def test_overlap_produces_coverage(self):
        """Chunks should have overlapping content (chunk N+1 starts near end of chunk N)."""
        text = "word " * 400
        chunks = self.chunker.chunk(text, {"doc_id": "x", "source_file": "x.txt"})
        if len(chunks) > 1:
            end_of_first = chunks[0].text[-30:]
            start_of_second = chunks[1].text[:100]
            # There should be some overlap
            assert any(word in start_of_second for word in end_of_first.split()[-3:])


class TestFileHandler:
    @pytest.fixture
    def handler(self, tmp_path):
        from cryptography.fernet import Fernet
        key = Fernet.generate_key()
        from src.ingestion.file_handler import FileHandler
        return FileHandler(
            upload_dir=str(tmp_path / "uploads"),
            processed_dir=str(tmp_path / "processed"),
            encryption_key=key,
        )

    @pytest.mark.asyncio
    async def test_upload_creates_encrypted_file(self, handler, tmp_path):
        from fastapi.datastructures import UploadFile
        content = b"Annual Report 2024 content"
        mock_file = MagicMock(spec=UploadFile)
        mock_file.filename = "annual_report.pdf"
        mock_file.read = AsyncMock(return_value=content)
        meta = await handler.handle_upload(mock_file, user="admin")
        assert meta["status"] == "uploaded"
        assert (Path(handler._upload_dir) / f"{meta['doc_id']}.enc").exists()

    @pytest.mark.asyncio
    async def test_reject_invalid_extension(self, handler):
        from fastapi.datastructures import UploadFile
        mock_file = MagicMock(spec=UploadFile)
        mock_file.filename = "malware.exe"
        mock_file.read = AsyncMock(return_value=b"bad data")
        with pytest.raises(HTTPException) as exc_info:
            await handler.handle_upload(mock_file, user="admin")
        assert exc_info.value.status_code == 415

    @pytest.mark.asyncio
    async def test_reject_oversized_file(self, handler):
        from fastapi.datastructures import UploadFile
        mock_file = MagicMock(spec=UploadFile)
        mock_file.filename = "huge.pdf"
        mock_file.read = AsyncMock(return_value=b"X" * (51 * 1024 * 1024))
        with pytest.raises(HTTPException) as exc_info:
            await handler.handle_upload(mock_file, user="admin", max_size_mb=50)
        assert exc_info.value.status_code in (413, 413)  # HTTP_413_CONTENT_TOO_LARGE

    def test_list_documents_empty(self, handler):
        assert handler.list_documents() == []

    @pytest.mark.asyncio
    async def test_delete_document(self, handler):
        from fastapi.datastructures import UploadFile
        content = b"test pdf content"
        mock_file = MagicMock(spec=UploadFile)
        mock_file.filename = "test.pdf"
        mock_file.read = AsyncMock(return_value=content)
        meta = await handler.handle_upload(mock_file, user="admin")
        doc_id = meta["doc_id"]
        result = handler.delete_document(doc_id)
        assert result is True
        assert handler.get_document(doc_id) is None
