"""
FinSight AI — Multi-Format Text Extractor
Supports PDF, DOCX, XLSX, CSV, and plain text. All extraction is in-memory.
"""
import csv
import io


class TextExtractor:
    """Extract plain text from file bytes based on file type."""

    def extract(self, file_bytes: bytes, file_type: str) -> str:
        """
        Extract text from file_bytes.
        file_type: 'pdf' | 'docx' | 'xlsx' | 'csv' | 'txt'
        """
        method = getattr(self, f"_extract_{file_type.lower()}", None)
        if method is None:
            raise ValueError(f"Unsupported file type: '{file_type}'")
        return method(file_bytes)

    def _extract_pdf(self, data: bytes) -> str:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data))
        return "\n\n".join(
            p.extract_text().strip() for p in reader.pages if p.extract_text()
        )

    def _extract_docx(self, data: bytes) -> str:
        import docx
        doc = docx.Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(c.text.strip() for c in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n\n".join(parts)

    def _extract_xlsx(self, data: bytes) -> str:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        sections = []
        for name in wb.sheetnames:
            rows = [
                "\t".join(str(v) if v is not None else "" for v in row)
                for row in wb[name].iter_rows(values_only=True)
                if any(v is not None for v in row)
            ]
            if rows:
                sections.append(f"[Sheet: {name}]\n" + "\n".join(rows))
        wb.close()
        return "\n\n".join(sections)

    def _extract_csv(self, data: bytes) -> str:
        text = data.decode("utf-8", errors="replace")
        rows = [
            "\t".join(row)
            for row in csv.reader(io.StringIO(text))
            if any(cell.strip() for cell in row)
        ]
        return "\n".join(rows)

    def _extract_txt(self, data: bytes) -> str:
        return data.decode("utf-8", errors="replace")
